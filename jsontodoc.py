import json
from docx import Document
from docx.shared import RGBColor, Inches

def hex_to_rgb(value):
    value = value.lstrip('#')
    lv = len(value)
    return tuple(int(value[i:i + lv // 3], 16) for i in range(0, lv, lv // 3))

def add_heading_and_content(doc, note):
    color = hex_to_rgb(note["bgColor"])
    rgb_color = RGBColor(*color)
    # Determine the heading level based on the color
    if note["color"] == 2:
        heading_level = 1
    elif note["color"] == 4:
        heading_level = 2
    elif note["color"] == 1:
        heading_level = 3
    elif note["color"] == 3:
        heading_level = 4
    else:
        heading_level = 0

    # Add the heading with the text from the "text" key
    heading = doc.add_heading(note["text"], heading_level)
    #set color of heading
    heading.runs[0].font.color.rgb = rgb_color
    # Add the content from the "note" key as a paragraph
    doc.add_paragraph(note["note"])

# Open the JSON file and read the data
with open("data.json", "r", encoding="utf-8") as json_file:
    data = json.load(json_file)

# Create a new Word document
doc = Document()

# Set the page size to A4 and add normal margins
section = doc.sections[0]
section.page_height = Inches(11.69)
section.page_width = Inches(8.27)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# Add a title to the document
doc.add_heading(data["bookName"], 0)

# Create a list of headings using the add_heading_and_content function
list(map(lambda note: add_heading_and_content(doc, note), data["noteList"]))

# Save the Word document
doc.save("document.docx")
